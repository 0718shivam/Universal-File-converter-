import pytesseract
from PIL import Image
import io
import os
import sys

# Try to set Tesseract path for Windows if needed
# Common Windows installation paths
if sys.platform == 'win32':
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe'.format(os.getenv('USERNAME', '')),
    ]
    
    tesseract_found = False
    for path in possible_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            tesseract_found = True
            break
    
    if not tesseract_found:
        # Try to find in PATH
        try:
            pytesseract.get_tesseract_version()
            tesseract_found = True
        except:
            pass
    
    if not tesseract_found:
        print("WARNING: Tesseract OCR not found. Please install from:")
        print("https://github.com/UB-Mannheim/tesseract/wiki")
        print("Or set the path manually in ocr_converter.py line 10")

try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("WARNING: python-docx not installed. DOCX export will not work.")

# Language mapping for Tesseract
LANGUAGE_MAP = {
    'english': 'eng',
    'hindi': 'hin',
    'sanskrit': 'san',
    'spanish': 'spa',
    'french': 'fra',
    'german': 'deu',
    'chinese': 'chi_sim',
    'japanese': 'jpn',
    'arabic': 'ara',
    'russian': 'rus',
    'portuguese': 'por',
    'italian': 'ita',
    'korean': 'kor',
    'tamil': 'tam',
    'telugu': 'tel',
    'bengali': 'ben',
    'gujarati': 'guj',
    'kannada': 'kan',
    'malayalam': 'mal',
    'marathi': 'mar',
    'punjabi': 'pan',
    'urdu': 'urd',
}

def extract_text_from_image(image_file, language='english'):
    """
    Extract text from image using OCR
    
    Args:
        image_file: File object containing the image
        language: Language name (e.g., 'english', 'hindi')
    
    Returns:
        Extracted text string
    """
    try:
        # Read image
        image_data = image_file.read()
        image_file.seek(0)  # Reset file pointer
        
        # Open image with PIL
        img = Image.open(io.BytesIO(image_data))
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Get Tesseract language code
        lang_code = LANGUAGE_MAP.get(language.lower(), 'eng')
        
        # Try to extract text with error handling
        try:
            # Use config for better accuracy
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(img, lang=lang_code, config=custom_config)
        except pytesseract.TesseractNotFoundError:
            raise Exception("Tesseract OCR is not installed. Please install Tesseract OCR from https://github.com/tesseract-ocr/tesseract")
        except pytesseract.TesseractError as e:
            # If language pack not found, try with English
            if 'Error' in str(e) or 'not found' in str(e).lower():
                if lang_code != 'eng':
                    text = pytesseract.image_to_string(img, lang='eng', config=custom_config)
                else:
                    raise Exception(f"OCR error: {str(e)}")
            else:
                raise Exception(f"OCR processing error: {str(e)}")
        
        return text.strip() if text else ""
    
    except Exception as e:
        error_msg = str(e)
        if 'Tesseract' in error_msg or 'tesseract' in error_msg.lower():
            raise Exception(f"Tesseract OCR error: {error_msg}. Please ensure Tesseract OCR is installed and in your system PATH.")
        raise Exception(f"OCR extraction failed: {error_msg}")

def create_text_file(text, filename='extracted_text.txt'):
    """Create a text file from extracted text"""
    text_buffer = io.BytesIO()
    text_buffer.write(text.encode('utf-8'))
    text_buffer.seek(0)
    return text_buffer

def create_word_file(text, filename='extracted_text.docx'):
    """Create a Word document from extracted text"""
    if not DOCX_AVAILABLE:
        raise Exception("python-docx is not installed. Please install it with: pip install python-docx")
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add text to document
    # Split by lines to preserve formatting
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph()  # Empty line
    
    # Save to buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def get_supported_languages():
    """Get list of supported languages"""
    return list(LANGUAGE_MAP.keys())

